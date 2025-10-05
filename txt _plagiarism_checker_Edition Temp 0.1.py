import re
import sys
from docx import Document  # 需安装python-docx库：pip install python-docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("""欢迎你使用适用于Windows,Linux,MacOS系统的汉语文本查重器!
      本文本查重器暂时只支持txt文本格式……
      只支持这几种系统的原因是导入文档时用的路径检查是否为路径用正则表达式来确认……
      本工具支持你查重的文档需要按照以下方式分割文章的组成:
      1.以空两格开头来分割这个段和上一个段。
      2.以无标点符号后的换行来分割诗歌的每一行。
      本程序将返回高亮的两个文件,内容是将两个原文件不同处高亮,并额外生成TXT和Word文件表明两个文件在第几行第几到第几的不同情况""")

# 初始化变量
f1, f2 = 0, 0
p = sys.platform

# 根据系统设置路径匹配正则
match p:
    case p.startswith('win'):
        matchstr = r'^[A-Za-z]:[\\/](?:[^*?"<>|:\n]+[\\/])*[^*?"<>|:\n]*\.txt$'
    case p.startswith('cygwin'):
        matchstr = r'^[A-Za-z]:[\\/](?:[^*?"<>|:\n]+[\\/])*[^*?"<>|:\n]*\.txt$'    
    case p.startswith('linux') | p.startswith('darwin'):  # darwin对应macOS
        matchstr = r'^(?:/[^/\0]+)*/[^/\0]*\.txt$'

# 获取并验证第一个文件路径
while f1 == 0:
    f1_temp = input("请输入第一个文档的绝对路径：")
    if re.fullmatch(matchstr, f1_temp):
        f1 = f1_temp
    else:
        print("路径格式错误，请重新输入（需为txt文件绝对路径）")

# 获取并验证第二个文件路径
while f2 == 0:
    f2_temp = input("请输入第二个文档的绝对路径：")
    if re.fullmatch(matchstr, f2_temp):
        f2 = f2_temp
    else:
        print("路径格式错误，请重新输入（需为txt文件绝对路径）")

# 读取文件内容
try:
    with open(f1, 'r', encoding='utf-8') as file:
        f1_lines = [line.rstrip('\n') for line in file.readlines()]
    with open(f2, 'r', encoding='utf-8') as file:
        f2_lines = [line.rstrip('\n') for line in file.readlines()]
except Exception as e:
    print(f"文件读取错误：{e}")
    sys.exit(1)

# 存储差异信息
diff_list = []

# 比对文件内容
max_lines = max(len(f1_lines), len(f2_lines))
for line_num in range(max_lines):
    # 获取当前行内容（超出范围视为空行）
    line1 = f1_lines[line_num] if line_num < len(f1_lines) else ''
    line2 = f2_lines[line_num] if line_num < len(f2_lines) else ''
    max_chars = max(len(line1), len(line2))
    i = 0
    
    while i < max_chars:
        # 处理一方超出长度的情况
        if i >= len(line1) or i >= len(line2):
            start = i
            end = max_chars - 1
            content1 = line1[start:] if i < len(line1) else ''
            content2 = line2[start:] if i < len(line2) else ''
            diff_list.append({
                'line': line_num + 1,  # 行号从1开始
                'f1_start': start + 1,
                'f1_end': start + len(content1),
                'f2_start': start + 1,
                'f2_end': start + len(content2),
                'f1_content': content1,
                'f2_content': content2
            })
            break
        
        # 发现差异字符
        if line1[i] != line2[i]:
            start = i
            # 查找连续差异的结束位置
            while i < max_chars and i < len(line1) and i < len(line2) and line1[i] != line2[i]:
                i += 1
            end = i - 1
            # 处理一方提前结束的情况
            while i < max_chars:
                if (i >= len(line1) or i >= len(line2)):
                    end = max_chars - 1
                    break
                if line1[i] == line2[i]:
                    break
                i += 1
                end = i - 1
            
            content1 = line1[start:end+1] if start < len(line1) else ''
            content2 = line2[start:end+1] if start < len(line2) else ''
            diff_list.append({
                'line': line_num + 1,
                'f1_start': start + 1,
                'f1_end': start + len(content1),
                'f2_start': start + 1,
                'f2_end': start + len(content2),
                'f1_content': content1,
                'f2_content': content2
            })
        i += 1

# 生成高亮文件（TXT格式）
def generate_highlighted_file(original_lines, diffs, output_filename):
    highlighted = []
    for line_num in range(len(original_lines)):
        line = original_lines[line_num]
        line_diffs = [d for d in diffs if d['line'] == line_num + 1]
        line_diffs.sort(key=lambda x: x['f1_start'])
        
        result = []
        prev_end = 0
        for d in line_diffs:
            start = d['f1_start'] - 1  # 转换为0索引
            end = d['f1_end'] - 1
            # 添加差异前的正常内容
            result.append(line[prev_end:start])
            # 添加高亮标记的差异内容
            result.append(f"【{line[start:end+1]}】")
            prev_end = end + 1
        # 添加剩余内容
        result.append(line[prev_end:])
        highlighted.append(''.join(result) + '\n')
    
    with open(output_filename, 'w', encoding='utf-8') as f:
        f.writelines(highlighted)

# 生成Word格式差异报告
def generate_word_report(diffs, f1_path, f2_path):
    doc = Document()
    
    # 设置标题
    title = doc.add_heading('文本差异报告', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 添加文件信息
    doc.add_paragraph(f"比对文件1：{f1_path}")
    doc.add_paragraph(f"比对文件2：{f2_path}")
    doc.add_paragraph("")  # 空行
    
    if not diffs:
        doc.add_paragraph("两个文件内容完全一致！")
    else:
        # 差异列表
        for idx, diff in enumerate(diffs, 1):
            # 差异编号
            p = doc.add_paragraph()
            run = p.add_run(f"差异 {idx}：")
            run.bold = True
            run.font.size = Pt(12)
            
            # 行号信息
            p = doc.add_paragraph(f"行号：第 {diff['line']} 行")
            p.paragraph_format.left_indent = Pt(15)
            
            # 文件1差异
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(15)
            p.add_run(f"文件1（第 {diff['f1_start']}-{diff['f1_end']} 字符）：").bold = True
            run = p.add_run(diff['f1_content'])
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色高亮
            
            # 文件2差异
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(15)
            p.add_run(f"文件2（第 {diff['f2_start']}-{diff['f2_end']} 字符）：").bold = True
            run = p.add_run(diff['f2_content'])
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)  # 蓝色高亮
            
            # 分隔线
            doc.add_paragraph("-" * 50)
    
    # 保存文档
    doc.save('差异报告.docx')

# 生成两个文件的高亮版本（TXT）
generate_highlighted_file(f1_lines, diff_list, '高亮_文件1.txt')
generate_highlighted_file(f2_lines, diff_list, '高亮_文件2.txt')

# 生成差异报告（TXT + Word）
with open('差异报告.txt', 'w', encoding='utf-8') as f:
    f.write("文本差异报告\n")
    f.write("==============\n\n")
    if not diff_list:
        f.write("两个文件内容完全一致！\n")
    else:
        for idx, diff in enumerate(diff_list, 1):
            f.write(f"差异 {idx}：\n")
            f.write(f"行号：第 {diff['line']} 行\n")
            f.write(f"文件1：第 {diff['f1_start']}-{diff['f1_end']} 字符：{diff['f1_content']}\n")
            f.write(f"文件2：第 {diff['f2_start']}-{diff['f2_end']} 字符：{diff['f2_content']}\n")
            f.write("---------\n")

# 生成Word报告
generate_word_report(diff_list, f1, f2)

print("比对完成！已生成：")
print("1. 高亮_文件1.txt")
print("2. 高亮_文件2.txt")
print("3. 差异报告.txt")
print("4. 差异报告.docx")
